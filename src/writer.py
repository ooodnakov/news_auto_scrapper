import os
import logging
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict

logger = logging.getLogger(__name__)

class ReportGenerator:
    def __init__(self, output_file="Report_Result.docx"):
        self.output_file = output_file
        self.doc = Document()

    def add_entry(self, entry: Dict):
        error_message = None
        if entry.get('status') == 'failed':
            error_message = entry.get('error') or 'Неизвестная ошибка'

        # 1. Original Metadata (Source, Date, Snippet, URL)
        source_line = entry.get('source') or 'Unknown Source'
        date_line = entry.get('date') or 'Unknown Date'
        snippet = entry.get('snippet') or entry.get('original_snippet') or ''
        url_line = entry.get('url') or ''

        self._add_preserved_paragraphs(source_line)
        self._add_preserved_paragraphs(date_line)
        if snippet:
            self._add_preserved_paragraphs(snippet)
        if url_line:
            paragraph = self.doc.add_paragraph(url_line)
            if paragraph.runs:
                paragraph.runs[0].font.color.rgb = RGBColor(0, 0, 255)
                paragraph.runs[0].underline = True

        self._add_preserved_paragraphs("▌")

        # 2. Screenshot
        telegram_path = entry.get('telegram_screenshot_path')
        if telegram_path and os.path.exists(telegram_path):
            self.doc.add_paragraph("Telegram App Screenshot:")
            try:
                self.doc.add_picture(telegram_path, width=Inches(6.0))
            except Exception as e:
                logger.error(f"Error inserting Telegram image {telegram_path}: {e}")
                self.doc.add_paragraph(f"[Error inserting Telegram image: {e}]")

        screenshot_path = entry.get('screenshot_path')
        if screenshot_path and os.path.exists(screenshot_path):
            self.doc.add_paragraph("Article Screenshot:")
            try:
                self.doc.add_picture(screenshot_path, width=Inches(6.0))
            except Exception as e:
                logger.error(f"Error inserting image {screenshot_path}: {e}")
                self.doc.add_paragraph(f"[Error inserting image: {e}]")
        
        # Spacing
        self.doc.add_paragraph("")

        # 3. Article Text
        content_blocks = entry.get('text_blocks')
        if content_blocks:
            for block in content_blocks:
                self._add_preserved_paragraphs(block)
        else:
            snippet = entry.get('snippet')
            if snippet:
                self._add_preserved_paragraphs(snippet)
            else:
                text_content = entry.get('full_text', '')
                self._add_preserved_paragraphs(text_content)
        
        # Error message if there was a failure
        if error_message:
            err_para = self.doc.add_paragraph(f"Ошибка обработки: {error_message}. Не удалось получить все данные.")
            err_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            err_para.runs[0].font.bold = True

        # Page Break between entries for cleanliness
        self.doc.add_page_break()

    def save(self):
        target_path = Path(self.output_file)
        try:
            self.doc.save(target_path)
            logger.info(f"💾 Report saved to {self.output_file}")
        except PermissionError as exc:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            fallback_name = f"Report_Result_{timestamp}.docx"
            fallback_path = target_path.parent / fallback_name if target_path.parent else Path(fallback_name)
            try:
                self.doc.save(fallback_path)
                logger.warning(
                    f"Permission denied writing {self.output_file} ({exc}). "
                    f"Saved report to fallback {fallback_path} instead."
                )
            except Exception as inner_exc:
                logger.error(f"Unable to save report to fallback location: {inner_exc}")

    def _add_preserved_paragraphs(self, text: str):
        """
        Adds paragraphs exactly as the text contains them so we keep the original layout.
        """
        if not text:
            return

        for line in text.split('\n'):
            paragraph = self.doc.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
